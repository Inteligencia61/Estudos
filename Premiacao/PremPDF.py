# -*- coding: utf-8 -*-
import pandas as pd
from datetime import datetime
import calendar

# Tenta importar gspread, mas permite execução local sem ele
try:
    import gspread
    from oauth2client.service_account import ServiceAccountCredentials
except ImportError:
    gspread = None
    ServiceAccountCredentials = None


# ==========================================================
# CONFIGURAÇÕES
# ==========================================================
MES_RELATORIO = 2
ANO_RELATORIO = 2026
TOP_N = 30

SHEET_ID_CONTRATOS = "1GEAy59-AjpkVuo3I1-ztt1adToDtEPMZrCGz_7B2byY"
SHEET_ID_BASE_INTELIGENCIA = "1v2Id3GE5HZkqq_2eEBt6WG6nVz3U0nnIYH6oS_v3gkA"

ABA_CONTRATOS = "Vendas"
ABA_DIM_CORRETOR = "Dim_Corretor"
ABA_DIM_GERENTE = "Dim_Gerente"

GERENTES_AC = {"HELIO JUNIO", "LUANA SALVINSKI"}
EXCECAO_OUTROS = {"FERNANDA LINDSAY"}  # sempre força PP, se quiser manter


# ==========================================================
# FUNÇÕES DE APOIO
# ==========================================================
def _to_float_br(x) -> float:
    if x is None:
        return 0.0
    if isinstance(x, (int, float)):
        return float(x)

    s = str(x).strip()
    if not s:
        return 0.0

    s = s.replace("R$", "").strip()

    if "," in s and "." in s:
        s = s.replace(".", "").replace(",", ".")
    elif "," in s and "." not in s:
        s = s.replace(",", ".")

    s = "".join(ch for ch in s if ch.isdigit() or ch in ".-")

    try:
        return float(s)
    except Exception:
        return 0.0


def limpar_nome(n):
    if pd.isna(n) or str(n).strip() in ["", "-", "nan", "NAN", "None"]:
        return ""
    return str(n).strip().upper()


def formatar_moeda_br(valor):
    try:
        return f"R$ {float(valor):,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
    except Exception:
        return "R$ 0,00"


def valor_positivo(x):
    if pd.isna(x):
        return False
    try:
        return _to_float_br(x) > 0
    except Exception:
        return False


def autenticar_google_sheets():
    if not gspread:
        raise ImportError("Biblioteca gspread não instalada.")

    scope = [
        "https://spreadsheets.google.com/feeds",
        "https://www.googleapis.com/auth/drive"
    ]
    creds = ServiceAccountCredentials.from_json_keyfile_name("../cred.json", scope)
    client = gspread.authorize(creds)
    return client


def carregar_aba_por_id(sheet_id, nome_aba):
    client = autenticar_google_sheets()
    planilha = client.open_by_key(sheet_id)
    aba = planilha.worksheet(nome_aba)
    dados = aba.get_all_values()

    if not dados:
        return pd.DataFrame()

    header = dados[0]
    rows = dados[1:]

    df = pd.DataFrame(rows, columns=header)
    df.columns = [str(c).strip() for c in df.columns]
    df = df.dropna(how="all")
    return df


# ==========================================================
# CARREGAMENTO DAS BASES
# ==========================================================
def carregar_contratos(sheet_id_contratos):
    return carregar_aba_por_id(sheet_id_contratos, ABA_CONTRATOS)


def carregar_base_inteligencia(sheet_id_base):
    df_corretor = carregar_aba_por_id(sheet_id_base, ABA_DIM_CORRETOR)
    df_gerente = carregar_aba_por_id(sheet_id_base, ABA_DIM_GERENTE)
    return df_corretor, df_gerente


# ==========================================================
# MAPA CORRETOR -> GERENTE
# ==========================================================
def montar_mapa_corretor_gerente(df_dim_corretor, df_dim_gerente):
    """
    Retorna:
    - mapa_nome_corretor_para_gerente_nome
    - mapa_nome_corretor_para_id_gerente
    """
    dfc = df_dim_corretor.copy()
    dfg = df_dim_gerente.copy()

    dfc.columns = [str(c).strip() for c in dfc.columns]
    dfg.columns = [str(c).strip() for c in dfg.columns]

    dfc["Nome_norm"] = dfc["Nome"].apply(limpar_nome)
    dfc["IdGerente_norm"] = dfc["IdGerente"].astype(str).str.strip().str.upper()

    dfg["IdGerente_norm"] = dfg["IdGerente"].astype(str).str.strip().str.upper()
    dfg["NomeGerente_norm"] = dfg["Nome"].apply(limpar_nome)

    mapa_idgerente_para_nome = dict(zip(dfg["IdGerente_norm"], dfg["NomeGerente_norm"]))

    mapa_corretor_para_idgerente = {}
    mapa_corretor_para_nomegerente = {}

    for _, row in dfc.iterrows():
        nome_corretor = row.get("Nome_norm", "")
        id_gerente = row.get("IdGerente_norm", "")

        if not nome_corretor:
            continue

        nome_gerente = mapa_idgerente_para_nome.get(id_gerente, "")
        mapa_corretor_para_idgerente[nome_corretor] = id_gerente
        mapa_corretor_para_nomegerente[nome_corretor] = nome_gerente

    return mapa_corretor_para_nomegerente, mapa_corretor_para_idgerente


# ==========================================================
# REGRAS DE TIME AC / PP
# ==========================================================
def bucket_time_do_corretor(nome_corretor, mapa_corretor_gerente):
    nome_corretor = limpar_nome(nome_corretor)

    if not nome_corretor:
        return "PP"

    if nome_corretor in EXCECAO_OUTROS:
        return "PP"

    gerente = limpar_nome(mapa_corretor_gerente.get(nome_corretor, ""))
    if "HELIO" in gerente or "LUANA" in gerente:
        return "AC"

    return "PP"


# ==========================================================
# EXTRAÇÃO DE PARTICIPANTES VÁLIDOS DA LINHA
# ==========================================================
def obter_vendedores_validos(row):
    vendedores = []

    if "$_Corretor_Venda_1" in row.index:
        if valor_positivo(row.get("$_Corretor_Venda_1")):
            nome = limpar_nome(row.get("Corretor_Venda_1_Nome"))
            if nome:
                vendedores.append(nome)
    else:
        nome = limpar_nome(row.get("Corretor_Venda_1_Nome"))
        if nome:
            vendedores.append(nome)

    if "$_Corretor_Venda_2" in row.index:
        if valor_positivo(row.get("$_Corretor_Venda_2")):
            nome = limpar_nome(row.get("Corretor_Venda_2_Nome"))
            if nome and nome not in vendedores:
                vendedores.append(nome)
    else:
        nome = limpar_nome(row.get("Corretor_Venda_2_Nome"))
        if nome and nome not in vendedores:
            vendedores.append(nome)

    return vendedores


def obter_captadores_validos(row):
    captadores = []

    if "$_Corretor_Captador_1" in row.index:
        if valor_positivo(row.get("$_Corretor_Captador_1")):
            nome = limpar_nome(row.get("Corretor_Captador_1_Nome"))
            if nome:
                captadores.append(nome)
    else:
        nome = limpar_nome(row.get("Corretor_Captador_1_Nome"))
        if nome:
            captadores.append(nome)

    if "$_Corretor_Captador_2" in row.index:
        if valor_positivo(row.get("$_Corretor_Captador_2")):
            nome = limpar_nome(row.get("Corretor_Captador_2_Nome"))
            if nome and nome not in captadores:
                captadores.append(nome)
    else:
        nome = limpar_nome(row.get("Corretor_Captador_2_Nome"))
        if nome and nome not in captadores:
            captadores.append(nome)

    return captadores


# ==========================================================
# AUXILIAR DE ACUMULAÇÃO
# ==========================================================
def _add(dic, chave, valor):
    if not chave:
        return
    dic[chave] = dic.get(chave, 0.0) + float(valor or 0.0)


def _init_rank_bloco():
    return {
        "VGV_CAP": {},
        "VGV_VEND": {},
        "VGV_GERAL": {},
        "VGC_CAP": {},
        "VGC_VEND": {},
        "VGC_GERAL": {},
    }


# ==========================================================
# RANKING DE CORRETORES
# ==========================================================
def processar_rankings_corretores_com_times(df_subset, mapa_corretor_gerente):
    """
    Regras corretas de VGV/VGC:
    - VGV venda: imóvel inteiro para cada vendedor da linha
    - VGV captação: imóvel inteiro para cada captador da linha
    - VGV geral: imóvel inteiro para cada participante da linha
    - VGC:
        * se tem venda e captação: 50% / 50%
        * se só venda: 100% venda
        * se só captação: 100% captação
        * cada lado divide entre seus participantes
    Separação AC/PP:
    - feita pelo gerente oficial do corretor na base mestre
    """
    res_ac = _init_rank_bloco()
    res_pp = _init_rank_bloco()
    res_geral = _init_rank_bloco()

    for _, row in df_subset.iterrows():
        v_imovel = _to_float_br(row.get("Valor_Negocio", 0))
        v_comissao_total = _to_float_br(row.get("Valor_Total_61", 0))

        vendedores = obter_vendedores_validos(row)
        captadores = obter_captadores_validos(row)

        set_vendedores = set(vendedores)
        set_captadores = set(captadores)
        todos = set_vendedores | set_captadores

        # -------------------------
        # VGV VENDA
        # -------------------------
        for nome in set_vendedores:
            bucket = bucket_time_do_corretor(nome, mapa_corretor_gerente)
            _add(res_geral["VGV_VEND"], nome, v_imovel)
            _add(res_geral["VGV_GERAL"], nome, v_imovel)

            if bucket == "AC":
                _add(res_ac["VGV_VEND"], nome, v_imovel)
                _add(res_ac["VGV_GERAL"], nome, v_imovel)
            else:
                _add(res_pp["VGV_VEND"], nome, v_imovel)
                _add(res_pp["VGV_GERAL"], nome, v_imovel)

        # -------------------------
        # VGV CAPTAÇÃO
        # -------------------------
        for nome in set_captadores:
            bucket = bucket_time_do_corretor(nome, mapa_corretor_gerente)
            _add(res_geral["VGV_CAP"], nome, v_imovel)

            if nome not in set_vendedores:
                _add(res_geral["VGV_GERAL"], nome, v_imovel)

            if bucket == "AC":
                _add(res_ac["VGV_CAP"], nome, v_imovel)
                if nome not in set_vendedores:
                    _add(res_ac["VGV_GERAL"], nome, v_imovel)
            else:
                _add(res_pp["VGV_CAP"], nome, v_imovel)
                if nome not in set_vendedores:
                    _add(res_pp["VGV_GERAL"], nome, v_imovel)

        # -------------------------
        # VGC
        # -------------------------
        tem_venda = len(set_vendedores) > 0
        tem_capt = len(set_captadores) > 0

        if tem_venda and tem_capt:
            parcela_venda = v_comissao_total * 0.5
            parcela_capt = v_comissao_total * 0.5
        elif tem_venda:
            parcela_venda = v_comissao_total
            parcela_capt = 0.0
        elif tem_capt:
            parcela_venda = 0.0
            parcela_capt = v_comissao_total
        else:
            parcela_venda = 0.0
            parcela_capt = 0.0

        # VGC venda
        if tem_venda and parcela_venda:
            por_vendedor = parcela_venda / len(set_vendedores)

            for nome in set_vendedores:
                bucket = bucket_time_do_corretor(nome, mapa_corretor_gerente)

                _add(res_geral["VGC_VEND"], nome, por_vendedor)
                _add(res_geral["VGC_GERAL"], nome, por_vendedor)

                if bucket == "AC":
                    _add(res_ac["VGC_VEND"], nome, por_vendedor)
                    _add(res_ac["VGC_GERAL"], nome, por_vendedor)
                else:
                    _add(res_pp["VGC_VEND"], nome, por_vendedor)
                    _add(res_pp["VGC_GERAL"], nome, por_vendedor)

        # VGC captação
        if tem_capt and parcela_capt:
            por_capt = parcela_capt / len(set_captadores)

            for nome in set_captadores:
                bucket = bucket_time_do_corretor(nome, mapa_corretor_gerente)

                _add(res_geral["VGC_CAP"], nome, por_capt)
                _add(res_geral["VGC_GERAL"], nome, por_capt)

                if bucket == "AC":
                    _add(res_ac["VGC_CAP"], nome, por_capt)
                    _add(res_ac["VGC_GERAL"], nome, por_capt)
                else:
                    _add(res_pp["VGC_CAP"], nome, por_capt)
                    _add(res_pp["VGC_GERAL"], nome, por_capt)

    return {
        "AC": res_ac,
        "PP": res_pp,
        "GERAL": res_geral,
    }


# ==========================================================
# RANKING DE GERENTES VIA BASE MESTRE
# ==========================================================
def processar_gerentes_via_dim_corretor(df_subset, mapa_corretor_gerente):
    vgv_cap, vgv_vend, vgv_geral = {}, {}, {}
    vgc_cap, vgc_vend, vgc_geral = {}, {}, {}

    detalhes_vgv_cap = {}
    detalhes_vgv_vend = {}
    detalhes_vgv_geral = {}

    for _, row in df_subset.iterrows():
        v_imovel = _to_float_br(row.get("Valor_Negocio", 0))
        v_comissao_total = _to_float_br(row.get("Valor_Total_61", 0))

        corretores_venda = obter_vendedores_validos(row)
        corretores_capt = obter_captadores_validos(row)

        gerentes_venda_dict = {}
        for corretor in corretores_venda:
            gerente = limpar_nome(mapa_corretor_gerente.get(corretor, ""))
            if gerente:
                gerentes_venda_dict.setdefault(gerente, []).append(corretor)

        gerentes_capt_dict = {}
        for corretor in corretores_capt:
            gerente = limpar_nome(mapa_corretor_gerente.get(corretor, ""))
            if gerente:
                gerentes_capt_dict.setdefault(gerente, []).append(corretor)

        # VGV VENDA
        for gerente, lista_corretores in gerentes_venda_dict.items():
            _add(vgv_vend, gerente, v_imovel)

            detalhes_vgv_vend.setdefault(gerente, {})
            quota = v_imovel / len(lista_corretores)

            for corretor in lista_corretores:
                _add(detalhes_vgv_vend[gerente], corretor, quota)

        # VGV CAPTAÇÃO
        for gerente, lista_corretores in gerentes_capt_dict.items():
            _add(vgv_cap, gerente, v_imovel)

            detalhes_vgv_cap.setdefault(gerente, {})
            quota = v_imovel / len(lista_corretores)

            for corretor in lista_corretores:
                _add(detalhes_vgv_cap[gerente], corretor, quota)

        # VGV GERAL
        gerentes_geral_dict = {}

        for gerente, lista_corretores in gerentes_venda_dict.items():
            gerentes_geral_dict.setdefault(gerente, set()).update(lista_corretores)

        for gerente, lista_corretores in gerentes_capt_dict.items():
            gerentes_geral_dict.setdefault(gerente, set()).update(lista_corretores)

        for gerente, conjunto_corretores in gerentes_geral_dict.items():
            _add(vgv_geral, gerente, v_imovel)

            detalhes_vgv_geral.setdefault(gerente, {})
            lista = list(conjunto_corretores)
            quota = v_imovel / len(lista)

            for corretor in lista:
                _add(detalhes_vgv_geral[gerente], corretor, quota)

        # VGC
        if gerentes_venda_dict and gerentes_capt_dict:
            parcela_venda = v_comissao_total * 0.5
            parcela_capt = v_comissao_total * 0.5
        elif gerentes_venda_dict:
            parcela_venda = v_comissao_total
            parcela_capt = 0.0
        elif gerentes_capt_dict:
            parcela_venda = 0.0
            parcela_capt = v_comissao_total
        else:
            parcela_venda = 0.0
            parcela_capt = 0.0

        if gerentes_venda_dict and parcela_venda:
            quota_gerente_venda = parcela_venda / len(gerentes_venda_dict)
            for gerente in gerentes_venda_dict:
                _add(vgc_vend, gerente, quota_gerente_venda)
                _add(vgc_geral, gerente, quota_gerente_venda)

        if gerentes_capt_dict and parcela_capt:
            quota_gerente_capt = parcela_capt / len(gerentes_capt_dict)
            for gerente in gerentes_capt_dict:
                _add(vgc_cap, gerente, quota_gerente_capt)
                _add(vgc_geral, gerente, quota_gerente_capt)

    return {
        "VGV_CAP": vgv_cap,
        "VGV_VEND": vgv_vend,
        "VGV_GERAL": vgv_geral,
        "VGC_CAP": vgc_cap,
        "VGC_VEND": vgc_vend,
        "VGC_GERAL": vgc_geral,
        "DETALHES_VGV_CAP": detalhes_vgv_cap,
        "DETALHES_VGV_VEND": detalhes_vgv_vend,
        "DETALHES_VGV_GERAL": detalhes_vgv_geral,
    }


# ==========================================================
# TABELAS
# ==========================================================
def ranking_dict_to_df(d: dict, top_n: int = 30) -> pd.DataFrame:
    if not d:
        return pd.DataFrame(columns=["Pos", "Nome", "Valor"])

    itens = sorted(d.items(), key=lambda x: x[1], reverse=True)[:top_n]
    df = pd.DataFrame(itens, columns=["Nome", "Valor"])
    df.insert(0, "Pos", range(1, len(df) + 1))
    return df


def montar_tabelas_relatorio(res_corretores, res_gerentes, top_n=30):
    tabelas = {}

    # CORRETORES AC
    tabelas["CORRETORES_AC_VGV_CAP"] = ranking_dict_to_df(res_corretores["AC"]["VGV_CAP"], top_n)
    tabelas["CORRETORES_AC_VGV_VEND"] = ranking_dict_to_df(res_corretores["AC"]["VGV_VEND"], top_n)
    tabelas["CORRETORES_AC_VGV_GERAL"] = ranking_dict_to_df(res_corretores["AC"]["VGV_GERAL"], top_n)

    tabelas["CORRETORES_AC_VGC_CAP"] = ranking_dict_to_df(res_corretores["AC"]["VGC_CAP"], top_n)
    tabelas["CORRETORES_AC_VGC_VEND"] = ranking_dict_to_df(res_corretores["AC"]["VGC_VEND"], top_n)
    tabelas["CORRETORES_AC_VGC_GERAL"] = ranking_dict_to_df(res_corretores["AC"]["VGC_GERAL"], top_n)

    # CORRETORES PP
    tabelas["CORRETORES_PP_VGV_CAP"] = ranking_dict_to_df(res_corretores["PP"]["VGV_CAP"], top_n)
    tabelas["CORRETORES_PP_VGV_VEND"] = ranking_dict_to_df(res_corretores["PP"]["VGV_VEND"], top_n)
    tabelas["CORRETORES_PP_VGV_GERAL"] = ranking_dict_to_df(res_corretores["PP"]["VGV_GERAL"], top_n)

    tabelas["CORRETORES_PP_VGC_CAP"] = ranking_dict_to_df(res_corretores["PP"]["VGC_CAP"], top_n)
    tabelas["CORRETORES_PP_VGC_VEND"] = ranking_dict_to_df(res_corretores["PP"]["VGC_VEND"], top_n)
    tabelas["CORRETORES_PP_VGC_GERAL"] = ranking_dict_to_df(res_corretores["PP"]["VGC_GERAL"], top_n)

    # CORRETORES GERAL
    tabelas["CORRETORES_GERAL_VGV_CAP"] = ranking_dict_to_df(res_corretores["GERAL"]["VGV_CAP"], top_n)
    tabelas["CORRETORES_GERAL_VGV_VEND"] = ranking_dict_to_df(res_corretores["GERAL"]["VGV_VEND"], top_n)
    tabelas["CORRETORES_GERAL_VGV_GERAL"] = ranking_dict_to_df(res_corretores["GERAL"]["VGV_GERAL"], top_n)

    tabelas["CORRETORES_GERAL_VGC_CAP"] = ranking_dict_to_df(res_corretores["GERAL"]["VGC_CAP"], top_n)
    tabelas["CORRETORES_GERAL_VGC_VEND"] = ranking_dict_to_df(res_corretores["GERAL"]["VGC_VEND"], top_n)
    tabelas["CORRETORES_GERAL_VGC_GERAL"] = ranking_dict_to_df(res_corretores["GERAL"]["VGC_GERAL"], top_n)

    # GERENTES
    tabelas["GERENTES_VGV_CAP"] = ranking_dict_to_df(res_gerentes["VGV_CAP"], top_n)
    tabelas["GERENTES_VGV_VEND"] = ranking_dict_to_df(res_gerentes["VGV_VEND"], top_n)
    tabelas["GERENTES_VGV_GERAL"] = ranking_dict_to_df(res_gerentes["VGV_GERAL"], top_n)

    tabelas["GERENTES_VGC_CAP"] = ranking_dict_to_df(res_gerentes["VGC_CAP"], top_n)
    tabelas["GERENTES_VGC_VEND"] = ranking_dict_to_df(res_gerentes["VGC_VEND"], top_n)
    tabelas["GERENTES_VGC_GERAL"] = ranking_dict_to_df(res_gerentes["VGC_GERAL"], top_n)

    return tabelas


# ==========================================================
# IMPRESSÃO CONSOLE
# ==========================================================
def imprimir_ranking(titulo, dados):
    print(f"\n--- {titulo} ---")
    if not dados:
        print("Sem dados.")
        return

    for nome, valor in sorted(dados.items(), key=lambda x: x[1], reverse=True):
        print(f"{nome.ljust(35)} | {formatar_moeda_br(valor)}")


# ==========================================================
# PDF GERAL
# ==========================================================
def gerar_pdf_ranking(tabelas: dict, caminho_pdf: str, titulo: str):
    from reportlab.lib.pagesizes import A4
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.lib import colors

    doc = SimpleDocTemplate(
        caminho_pdf,
        pagesize=A4,
        rightMargin=30, leftMargin=30,
        topMargin=30, bottomMargin=30
    )

    styles = getSampleStyleSheet()
    story = []

    story.append(Paragraph(titulo, styles["Title"]))
    story.append(Paragraph(f"Gerado em: {datetime.now().strftime('%d/%m/%Y %H:%M')}", styles["Normal"]))
    story.append(Spacer(1, 12))

    def add_table(df: pd.DataFrame, secao: str):
        story.append(Paragraph(secao.replace("_", " "), styles["Heading2"]))

        if df.empty:
            story.append(Paragraph("Sem dados.", styles["Normal"]))
            story.append(Spacer(1, 10))
            return

        df2 = df.copy()
        df2["Valor"] = df2["Valor"].apply(formatar_moeda_br)

        data = [df2.columns.tolist()] + df2.values.tolist()

        t = Table(data, hAlign="LEFT", colWidths=[40, 290, 120])
        t.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1f2937")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("GRID", (0, 0), (-1, -1), 0.3, colors.grey),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.whitesmoke, colors.lightgrey]),
            ("FONTSIZE", (0, 0), (-1, -1), 9),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ]))
        story.append(t)
        story.append(Spacer(1, 14))

    for i, (secao, df) in enumerate(tabelas.items(), start=1):
        add_table(df, secao)
        if i % 3 == 0 and i < len(tabelas):
            story.append(PageBreak())

    doc.build(story)


# ==========================================================
# DOCX GERAL
# ==========================================================
def gerar_docx_ranking(tabelas: dict, caminho_docx: str, titulo: str):
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    h = doc.add_heading(titulo, level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph(f"Gerado em: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
    p.runs[0].font.size = Pt(10)

    doc.add_paragraph("")

    for secao, df in tabelas.items():
        doc.add_heading(secao.replace("_", " "), level=2)

        if df.empty:
            doc.add_paragraph("Sem dados.")
            doc.add_paragraph("")
            continue

        table = doc.add_table(rows=1, cols=3)
        hdr = table.rows[0].cells
        hdr[0].text = "Pos"
        hdr[1].text = "Nome"
        hdr[2].text = "Valor"

        for _, r in df.iterrows():
            row = table.add_row().cells
            row[0].text = str(int(r["Pos"]))
            row[1].text = str(r["Nome"])
            row[2].text = formatar_moeda_br(r["Valor"])

        doc.add_paragraph("")

    doc.save(caminho_docx)


# ==========================================================
# PDF DETALHADO DE GERENTES
# ==========================================================
def gerar_pdf_gerentes_detalhado(res_gerentes: dict, caminho_pdf: str, titulo: str, top_n: int = 30):
    from reportlab.lib.pagesizes import A4
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.lib import colors

    doc = SimpleDocTemplate(
        caminho_pdf,
        pagesize=A4,
        rightMargin=30, leftMargin=30,
        topMargin=30, bottomMargin=30
    )

    styles = getSampleStyleSheet()
    story = []

    story.append(Paragraph(titulo, styles["Title"]))
    story.append(Paragraph(f"Gerado em: {datetime.now().strftime('%d/%m/%Y %H:%M')}", styles["Normal"]))
    story.append(Spacer(1, 12))

    ranking_gerentes = sorted(
        res_gerentes["VGV_GERAL"].items(),
        key=lambda x: x[1],
        reverse=True
    )[:top_n]

    story.append(Paragraph("Ranking de Gerentes - VGV Geral", styles["Heading2"]))

    if not ranking_gerentes:
        story.append(Paragraph("Sem dados.", styles["Normal"]))
        doc.build(story)
        return

    resumo = [["Pos", "Gerente", "VGV Geral"]]
    for i, (gerente, valor) in enumerate(ranking_gerentes, start=1):
        resumo.append([str(i), gerente, formatar_moeda_br(valor)])

    t = Table(resumo, hAlign="LEFT", colWidths=[35, 290, 120])
    t.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1f2937")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("GRID", (0, 0), (-1, -1), 0.3, colors.grey),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.whitesmoke, colors.lightgrey]),
        ("FONTSIZE", (0, 0), (-1, -1), 9),
    ]))
    story.append(t)
    story.append(Spacer(1, 16))

    detalhes = res_gerentes.get("DETALHES_VGV_GERAL", {})

    for i, (gerente, valor_total) in enumerate(ranking_gerentes, start=1):
        story.append(Paragraph(f"{i}. {gerente} - {formatar_moeda_br(valor_total)}", styles["Heading2"]))
        story.append(Spacer(1, 6))

        corretores = detalhes.get(gerente, {})
        corretores_ordenados = sorted(corretores.items(), key=lambda x: x[1], reverse=True)

        data = [["Pos", "Corretor", "VGV usado na soma"]]
        soma = 0.0

        for j, (corretor, valor) in enumerate(corretores_ordenados, start=1):
            soma += float(valor or 0.0)
            data.append([str(j), corretor, formatar_moeda_br(valor)])

        data.append(["", "TOTAL DETALHADO", formatar_moeda_br(soma)])

        td = Table(data, hAlign="LEFT", colWidths=[35, 290, 120])
        td.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#374151")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("GRID", (0, 0), (-1, -1), 0.25, colors.grey),
            ("ROWBACKGROUNDS", (0, 1), (-1, -2), [colors.beige, colors.whitesmoke]),
            ("BACKGROUND", (0, -1), (-1, -1), colors.HexColor("#d1d5db")),
            ("FONTNAME", (0, -1), (-1, -1), "Helvetica-Bold"),
            ("FONTSIZE", (0, 0), (-1, -1), 8.5),
        ]))
        story.append(td)
        story.append(Spacer(1, 14))

        if i < len(ranking_gerentes):
            story.append(PageBreak())

    doc.build(story)


# ==========================================================
# MAIN
# ==========================================================
def main(
    sheet_id_contratos=SHEET_ID_CONTRATOS,
    sheet_id_base=SHEET_ID_BASE_INTELIGENCIA,
    usar_ano_todo=False
):
    # 1) Carrega dados
    df = carregar_contratos(sheet_id_contratos)
    df_dim_corretor, df_dim_gerente = carregar_base_inteligencia(sheet_id_base)

    # 2) Converte financeiros
    for col in ["Valor_Negocio", "Valor_Total_61"]:
        if col in df.columns:
            df[col] = df[col].apply(_to_float_br)

    # 3) Filtro de data
    df["Data_Contrato"] = pd.to_datetime(df["Data_Contrato"], errors="coerce")

    if usar_ano_todo:
        df_mes = df[
            (df["Data_Contrato"] >= f"{ANO_RELATORIO}-01-01") &
            (df["Data_Contrato"] <= f"{ANO_RELATORIO}-12-31")
        ].copy()
    else:
        ultimo_dia = calendar.monthrange(ANO_RELATORIO, MES_RELATORIO)[1]
        df_mes = df[
            (df["Data_Contrato"] >= f"{ANO_RELATORIO}-{MES_RELATORIO:02d}-01") &
            (df["Data_Contrato"] <= f"{ANO_RELATORIO}-{MES_RELATORIO:02d}-{ultimo_dia:02d}")
        ].copy()

    # 4) Mapa oficial corretor -> gerente
    mapa_corretor_gerente, _ = montar_mapa_corretor_gerente(df_dim_corretor, df_dim_gerente)

    # 5) Processa corretores com separação AC/PP correta
    res_corretores = processar_rankings_corretores_com_times(df_mes, mapa_corretor_gerente)

    # 6) Processa gerentes via base mestre
    res_gerentes = processar_gerentes_via_dim_corretor(df_mes, mapa_corretor_gerente)

    # 7) Console
    print("\n" + "=" * 80)
    print(f"RANKING DE CORRETORES - {MES_RELATORIO:02d}/{ANO_RELATORIO}")
    print("=" * 80)

    print("\n>>> CORRETORES AC - VGV")
    imprimir_ranking("AC | VGV CAP", res_corretores["AC"]["VGV_CAP"])
    imprimir_ranking("AC | VGV VEND", res_corretores["AC"]["VGV_VEND"])
    imprimir_ranking("AC | VGV GERAL", res_corretores["AC"]["VGV_GERAL"])

    print("\n>>> CORRETORES AC - VGC")
    imprimir_ranking("AC | VGC CAP", res_corretores["AC"]["VGC_CAP"])
    imprimir_ranking("AC | VGC VEND", res_corretores["AC"]["VGC_VEND"])
    imprimir_ranking("AC | VGC GERAL", res_corretores["AC"]["VGC_GERAL"])

    print("\n>>> CORRETORES PP - VGV")
    imprimir_ranking("PP | VGV CAP", res_corretores["PP"]["VGV_CAP"])
    imprimir_ranking("PP | VGV VEND", res_corretores["PP"]["VGV_VEND"])
    imprimir_ranking("PP | VGV GERAL", res_corretores["PP"]["VGV_GERAL"])

    print("\n>>> CORRETORES PP - VGC")
    imprimir_ranking("PP | VGC CAP", res_corretores["PP"]["VGC_CAP"])
    imprimir_ranking("PP | VGC VEND", res_corretores["PP"]["VGC_VEND"])
    imprimir_ranking("PP | VGC GERAL", res_corretores["PP"]["VGC_GERAL"])

    print("\n>>> CORRETORES GERAL - VGV")
    imprimir_ranking("GERAL | VGV CAP", res_corretores["GERAL"]["VGV_CAP"])
    imprimir_ranking("GERAL | VGV VEND", res_corretores["GERAL"]["VGV_VEND"])
    imprimir_ranking("GERAL | VGV GERAL", res_corretores["GERAL"]["VGV_GERAL"])

    print("\n>>> CORRETORES GERAL - VGC")
    imprimir_ranking("GERAL | VGC CAP", res_corretores["GERAL"]["VGC_CAP"])
    imprimir_ranking("GERAL | VGC VEND", res_corretores["GERAL"]["VGC_VEND"])
    imprimir_ranking("GERAL | VGC GERAL", res_corretores["GERAL"]["VGC_GERAL"])

    print("\n>>> GERENTES - VGV")
    imprimir_ranking("GERENTES | VGV CAP", res_gerentes["VGV_CAP"])
    imprimir_ranking("GERENTES | VGV VEND", res_gerentes["VGV_VEND"])
    imprimir_ranking("GERENTES | VGV GERAL", res_gerentes["VGV_GERAL"])

    print("\n>>> GERENTES - VGC")
    imprimir_ranking("GERENTES | VGC CAP", res_gerentes["VGC_CAP"])
    imprimir_ranking("GERENTES | VGC VEND", res_gerentes["VGC_VEND"])
    imprimir_ranking("GERENTES | VGC GERAL", res_gerentes["VGC_GERAL"])

    # 8) Gera arquivos
    sufixo_periodo = f"{ANO_RELATORIO}_{MES_RELATORIO:02d}" if not usar_ano_todo else f"{ANO_RELATORIO}_ANO_TODO"
    titulo = f"Relatório de Rankings - {MES_RELATORIO:02d}/{ANO_RELATORIO}" if not usar_ano_todo else f"Relatório de Rankings - Ano {ANO_RELATORIO}"

    tabelas = montar_tabelas_relatorio(res_corretores, res_gerentes, top_n=TOP_N)

    nome_pdf = f"ranking_completo_{sufixo_periodo}.pdf"
    nome_docx = f"ranking_completo_{sufixo_periodo}.docx"
    nome_pdf_gerentes = f"ranking_gerentes_detalhado_{sufixo_periodo}.pdf"

    gerar_pdf_ranking(tabelas, nome_pdf, titulo)
    gerar_docx_ranking(tabelas, nome_docx, titulo)

    gerar_pdf_gerentes_detalhado(
        res_gerentes,
        nome_pdf_gerentes,
        f"Ranking de Gerentes com base na Dim_Corretor - {titulo}",
        top_n=TOP_N
    )

    print("\nArquivos gerados:")
    print(f"- {nome_pdf}")
    print(f"- {nome_docx}")
    print(f"- {nome_pdf_gerentes}")


if __name__ == "__main__":
    main()